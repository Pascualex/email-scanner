import re
import os
import ctypes
import fitz
from io import StringIO
from docx import Document
from docx.shared import Pt
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT


txt_name = 'usuarios'
txt_final_name = None
docx_name = 'tabla_usuarios'
docx_final_name = None

users = []

name = None
lastname = None
dni = None
email = None
occupation = None


def read_pdf():
    try:
        doc = fitz.open('correos.pdf')
        for i in range(doc.pageCount):
            p = doc.loadPage(i)
            text = p.getText()

            print('\n=======================> Página', i, '\n')
            print(text)

            for line in text.splitlines():
                parts = re.compile('[ ]*:[ ]*').split(line)
                if len(parts) == 2:
                    process_field(parts[0], parts[1])

        store_fields()

        return True
    except:        
        title = 'Error'
        message = 'No se ha encontrado el fichero correos.pdf. Por favor, colóquelo en la misma carpeta que este ejecutable.'
        ctypes.windll.user32.MessageBoxW(0, message, title, 0x40000)

        return False


def process_field(field, content):
    global name, lastname, dni, email, occupation
    if field == 'De':
        store_fields()
        reset_fields()
    elif field == 'Nombre':
        name = content
    elif field == 'Apellidos':
        lastname = content
    elif field == 'DNI':
        dni = content
    elif field == 'Correo electrónico':
        email = content
    elif field == 'Puesto de trabajo':
        occupation = content


def reset_fields():
    global name, lastname, dni, email, occupation
    name = None
    lastname = None
    dni = None
    email = None
    occupation = None


def store_fields():
    if name is None or lastname is None or dni is None or email is None or occupation is None:
        return
    users.append({
        'name': name, 
        'lastname': lastname,
        'dni': dni, 
        'email': email,
        'occupation': occupation
    })


def write_txt():
    global txt_final_name

    if os.path.isfile('./' + txt_name + '.txt') is True:
        file_id = 1
        while os.path.isfile('./' + txt_name + '_' + str(file_id) + '.txt') is True:
            file_id += 1
        txt_final_name = txt_name + '_' + str(file_id) + '.txt'
    else:
        txt_final_name = txt_name + '.txt'
        
    file = open(txt_final_name, 'w')
    file.write('nombre|apellidos|dni|correo|puesto\n')

    for user in users:
        file.write(user['name'] + '|' + user['lastname'] + '|' + user['dni'] + '|' + user['email'] + '|' + user['occupation'] + '\n')

    file.close()


def write_word():
    global docx_final_name

    document = Document()
    style = document.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    section = document.sections[-1]
    new_width = section.page_height
    new_height = section.page_width
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = new_width
    section.page_height = new_height

    table = document.add_table(rows=1, cols=6)
    table.style = 'Table Grid'

    cells = table.rows[0].cells
    cells[0].text = 'Nº'
    cells[0].paragraphs[0].runs[0].font.bold = True
    cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cells[1].text = 'NOMBRE'
    cells[1].paragraphs[0].runs[0].font.bold = True
    cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cells[2].text = 'APELLIDOS'
    cells[2].paragraphs[0].runs[0].font.bold = True
    cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cells[3].text = 'NIF'
    cells[3].paragraphs[0].runs[0].font.bold = True
    cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cells[4].text = 'CORREO'
    cells[4].paragraphs[0].runs[0].font.bold = True
    cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cells[5].text = 'OCUPACIÓN ACTUAL'
    cells[5].paragraphs[0].runs[0].font.bold = True
    cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    position = 1
    for user in users:
        cells = table.add_row().cells
        cells[0].text = str(position)
        cells[1].text = user['name']
        cells[2].text = user['lastname']
        cells[3].text = user['dni']
        cells[4].text = user['email']
        cells[5].text = user['occupation']
        
        position += 1

    widths = (Cm(0.85), Cm(3.25), Cm(4.39), Cm(2.25), Cm(5.25), Cm(9))
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width

    if os.path.isfile(docx_name + '.docx') is True:
        file_id = 1
        while os.path.isfile(docx_name + '_' + str(file_id) + '.docx') is True:
            file_id += 1
        docx_final_name = docx_name + '_' + str(file_id) + '.docx'
    else:
        docx_final_name = docx_name + '.docx'
    
    document.save(docx_final_name)


if __name__ == '__main__':
    if read_pdf():
        write_txt()
        write_word()
        title = 'Éxito'
        message = 'Se han procesado correctamente ' + str(len(users)) + ' usuarios. Los resultados se han almacenado en ' + txt_final_name + ' y ' + docx_final_name + '.'
        ctypes.windll.user32.MessageBoxW(0, message, title, 0x40000)